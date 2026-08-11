# -*- coding: utf-8 -*-
"""
별지 제34호 서식(통합신청서) 자동 작성 모듈
- {{TAG}} 치환 방식이 아니라 '셀 좌표(row, col) 직접 기입' 방식
- 이유: 법무부 원본 서식에는 태그가 없고, 값이 들어갈 자리가 전부 '빈 병합 셀'임
"""
import io
import os
import re
import docx
from docx.shared import Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH

TEMPLATE_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "template.docx")
FONT_NAME = "돋움"
FONT_SIZE = Pt(9)

# ── 서식 좌표 지도 (table 0 기준, 실측값) ────────────────────────────
FIELD_MAP = {
    "surname":        (14, 2),    # 성 Surname
    "given_names":    (14, 12),   # 명 Given names
    "hanja_name":     (14, 34),   # 漢字姓名
    "birth_yyyy":     (16, 3),
    "birth_mm":       (16, 13),
    "birth_dd":       (16, 17),
    "nationality":    (15, 38),   # 국적 (r15~r17 병합)
    "passport_no":    (18, 2),
    "passport_issue": (18, 22),
    "passport_expiry":(18, 38),
    "address_kr":     (19, 2),    # 대한민국 내 주소
    "tel":            (20, 6),
    "cell_phone":     (20, 34),
    "address_home":   (21, 6),    # 본국 주소
    "home_phone":     (21, 39),
    "school_name":    (22, 27),   # 학교 이름
    "school_phone":   (22, 39),
    "email":          (27, 27),   # 전자우편
    "apply_date":     (29, 9),    # 신청일
    "signature":      (29, 35),   # 신청인 서명
    "consent_sign":   (35, 1),    # 행정정보 공동이용 동의서 서명란
}

# 외국인등록번호 13칸 (한 칸에 한 글자)
ARC_CELLS = [(17, c) for c in
             (7, 9, 13, 14, 17, 19, 22, 24, 26, 27, 29, 30, 33)]

# 신청 종류 체크박스
CHECKBOX_MAP = {
    "외국인등록":       (4, 0),
    "체류자격외활동허가": (4, 12),
    "등록증재발급":     (6, 0),
    "근무처변경":       (6, 12),
    "체류기간연장허가":  (7, 0),   # ← 우리 서비스의 기본값
    "재입국허가":       (7, 12),
    "체류자격변경허가":  (8, 0),
    "체류지변경신고":    (8, 12),
    "체류자격부여":      (10, 0),
    "등록사항변경신고":  (10, 12),
}
SEX_CELL = (15, 27)   # '[ ]남 M\n[ ]여 F'

# 주소·이메일처럼 긴 값은 좌측 정렬
LEFT_ALIGN_FIELDS = {"address_kr", "address_home", "email"}

# 한 셀 안에 체크박스가 여러 개 들어있는 항목 (라벨로 지정해서 체크)
INLINE_CHECK_CELLS = {
    "school_status": (22, 2),   # 미취학[] 초[] 중[] 고[]
    "school_type":   (23, 9),   # 교육청 인가[] / 대안학교[]
}

PHOTO_CELL = (4, 34)


def _write(cell, text, align=WD_ALIGN_PARAGRAPH.CENTER):
    """빈 셀에 값 기입. 셀 서식(테두리/배경)은 건드리지 않고 run만 추가."""
    if text is None or text == "":
        return
    p = cell.paragraphs[0]
    for r in list(p.runs):          # 기존 run 제거 (빈 셀이면 없음)
        r._element.getparent().remove(r._element)
    run = p.add_run(str(text))
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE
    # 한글 폰트는 eastAsia 속성까지 지정해야 Word에서 제대로 적용됨
    run._element.rPr.rFonts.set(
        docx.oxml.ns.qn("w:eastAsia"), FONT_NAME)
    p.alignment = align


def _tick(cell, label=None):
    """'[  ]' 를 '[V]' 로 바꿔 체크 표시. label이 있으면 그 줄만 체크."""
    for p in cell.paragraphs:
        if label and label not in p.text:
            continue
        for run in p.runs:
            if "[" in run.text and "]" in run.text:
                run.text = run.text.replace("[  ]", "[V]").replace("[ ]", "[V]")
                run.font.bold = True
                if label:
                    return


def _tick_inline(cell, label):
    """'미취학[]', '초[ ]'처럼 한 셀에 여러 체크박스가 있을 때 특정 항목만 체크.
    Word가 텍스트를 여러 run으로 쪼개 놓으므로 문단 단위로 다시 쓴다."""
    if not label:
        return
    for p in cell.paragraphs:
        text = "".join(r.text for r in p.runs)
        # '라벨' 뒤의 대괄호 안을 V로 (사이 공백 허용)
        new_text, n = re.subn(
            re.escape(label) + r"(\s*)\[\s*\]",
            lambda m: f"{label}{m.group(1)}[V]",
            text,
            count=1,
        )
        if n:
            keep = p.runs[0]
            keep.text = new_text
            keep.font.bold = True
            for r in p.runs[1:]:
                r._element.getparent().remove(r._element)
            return


def _insert_photo(cell, photo_bytes, width_mm=32, height_mm=41):
    """여권용 사진을 사진란에 삽입. 외국인등록·등록증 재발급 신청 시에만 사용."""
    for p in list(cell.paragraphs):
        p._element.getparent().remove(p._element)
    new_p = cell.add_paragraph()
    new_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    new_p.add_run().add_picture(
        io.BytesIO(photo_bytes), width=Mm(width_mm), height=Mm(height_mm)
    )


def _split_name(full_name: str):
    """Vision Agent가 'HONG SAMPLE' 처럼 통으로 준 이름을 성/명으로 분리.
    외국인등록증 영문명은 보통 'SURNAME GIVEN NAMES' 순서."""
    if not full_name:
        return "", ""
    parts = full_name.strip().split()
    if len(parts) == 1:
        return parts[0], ""
    return parts[0], " ".join(parts[1:])


def fill_application(user_data: dict, template_path: str = TEMPLATE_PATH,
                     photo_bytes: bytes = None) -> bytes:
    doc = docx.Document(template_path)
    t = doc.tables[0]

    def cell(rc):
        return t.rows[rc[0]].cells[rc[1]]

    data = dict(user_data or {})

    # 이름이 통으로만 있으면 성/명 분리
    if not data.get("surname"):
        sn, gn = _split_name(data.get("name", ""))
        data.setdefault("surname", sn)
        data.setdefault("given_names", gn)

    # 1) FIELD_MAP 전체를 순회하며 값이 있는 칸만 채운다
    for key, rc in FIELD_MAP.items():
        val = data.get(key)
        if not val:
            continue
        align = (WD_ALIGN_PARAGRAPH.LEFT if key in LEFT_ALIGN_FIELDS
                 else WD_ALIGN_PARAGRAPH.CENTER)
        _write(cell(rc), val, align)

    # 2) 외국인등록번호 13칸 분해 기입
    arc = "".join(ch for ch in str(data.get("arc_no", "")) if ch.isdigit())
    for i, rc in enumerate(ARC_CELLS):
        if i < len(arc):
            _write(cell(rc), arc[i])

    # 3) 신청 종류 체크 (기본: 체류기간 연장허가)
    kind = data.get("application_kind", "체류기간연장허가")
    if kind in CHECKBOX_MAP:
        _tick(cell(CHECKBOX_MAP[kind]))

    # 4) 성별 체크
    sex = str(data.get("sex") or "").upper()
    if sex.startswith("M") or sex == "남":
        _tick(cell(SEX_CELL), label="남")
    elif sex.startswith("F") or sex == "여":
        _tick(cell(SEX_CELL), label="여")

    # 5) 재학 여부 / 학교 종류 (초·중·고 유학생만 해당. 대학생은 공란)
    #    허용값: 미취학 / 초 / 중 / 고
    _tick_inline(cell(INLINE_CHECK_CELLS["school_status"]),
                 data.get("school_status"))
    #    허용값: 교육청 인가 / 대안학교  (원본 서식상 '교육청 비인가'에는 칸이 없음)
    _tick_inline(cell(INLINE_CHECK_CELLS["school_type"]),
                 data.get("school_type"))

    # 6) 여권용 사진 (외국인등록·등록증 재발급 신청 시에만)
    if photo_bytes:
        _insert_photo(cell(PHOTO_CELL), photo_bytes)

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()


if __name__ == "__main__":
    sample = {
        "name": "HONG SAMPLE",
        "nationality": "VIETNAM",
        "visa_type": "D-2-6",
        "arc_no": "010203-5123456",
        "sex": "M",
        "birth_yyyy": "2001", "birth_mm": "02", "birth_dd": "03",
        "passport_no": "M12345678",
        "passport_issue": "2023.05.01",
        "passport_expiry": "2033.04.30",
        "address_kr": "경기도 수원시 영통구 월드컵로 206 기숙사 A동 512호",
        "cell_phone": "010-1234-5678",
        "email": "hong.sample@example.ac.kr",
        "school_name": "아주대학교",
        "apply_date": "2026. 08. 10.",
    }
    open("out.docx", "wb").write(fill_application(sample))
    print("written out.docx")